import docx
from docx.shared import Mm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_LINE_SPACING
from datetime import datetime, date, time
import calendar
import Google_docs

report_month = datetime.timetuple(datetime.now())[1] - 1

current_year = datetime.timetuple(datetime.now())[0]
current_month = datetime.timetuple(datetime.now())[1]
corrent_day = datetime.timetuple(datetime.now())[2]

quantity_days_of_previous_month = calendar.monthrange(current_year, report_month)[1]

doc = docx.Document()
# Стиль текста по умолчанию
style = doc.styles['Normal'] # Стиль для всего документа
style.font.name = 'TimesNewRomanPSMT' # Шрифт для всего документа
style.font.size = Pt(11) # Размер шрифта

# Добавление дотабличного текста
header = doc.add_paragraph('')
header.add_run('Акт сдачи-приемки оказанных услуг').bold = True
header.alignment = WD_ALIGN_PARAGRAPH.CENTER

par1 = doc.add_paragraph('')
par1.add_run('Акционерное общество «МР Групп» ').bold = True
par1.add_run('(АО «МР Групп»), зарегистрированное и действующее в соответствии с законодательством '
             'Российской Федерации, именуемое в дальнейшем «Заказчик», в лице Исполнительного директора '
             'Обухова Александра Анатольевича, действующего на основании доверенности '
             'от 31.12.2021 No ИД-2022, с одной стороны, и ')
par1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
fmt = par1.paragraph_format
fmt.first_line_indent = Mm(12.7)

par2 = doc.add_paragraph('')
par2.add_run("Общество с ограниченной ответственностью «Мультиспейс Павелецкая» ").bold = True
par2.add_run('(ООО «Мультиспейс Павелецкая»), именуемое в дальнейшем «Исполнитель», '
             'в лице Генерального директора управляющей организации ООО «Прайдекс Констракшн» '
             'Бренева Кирилла Сергеевича, действующего на основании Устава и Договора No 1 передачи '
             'полномочий единоличного исполнительного органа ООО «Электросистемс» от 18.11.2019г., с '
             'другой стороны, далее совместно именуемые «Стороны», а по отдельности – «Сторона», '
             'составили настоящий акт (далее - Акт) о нижеследующем:')
par2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
fmt2 = par2.paragraph_format
fmt2.first_line_indent = Mm(12.7)

par3 = doc.add_paragraph(f'1. Исполнителем были оказаны следующие услуги за период с 01.{report_month:02}.{current_year}г. по {quantity_days_of_previous_month:02}.{report_month:02}.{current_year}г. '
                         '(далее - Отчетный период) по Договору возмездного оказания услуг '
                         'No 41254 (далее - Договор): ')
par3.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
fmt3 = par3.paragraph_format
fmt3.first_line_indent = Mm(12.7)

review = Google_docs.last_month_review

table = doc.add_table(len(review) + 1, 3) # (строка, колонка) | Создаём таблицу
table.style = 'Table Grid' # Стиль таблицы (Шоб рамки были видны)
cell = table.cell(0, 0).add_paragraph('')
cell.add_run("Дата")
cell.alignment = WD_ALIGN_PARAGRAPH.CENTER
cell = table.cell(0, 1).add_paragraph('')
cell.add_run("Описание услуги/инцидента")
cell.alignment = WD_ALIGN_PARAGRAPH.CENTER
cell = table.cell(0, 2).add_paragraph('')
cell.add_run("Актуальный статус")
cell.alignment = WD_ALIGN_PARAGRAPH.CENTER
for row in range(len(review)):
    for column in range(3):
        cell = table.cell(row + 1, column).add_paragraph('')
        cell.add_run(review[row][column])
        cell.alignment = WD_ALIGN_PARAGRAPH.CENTER

deleter = doc.add_paragraph('\n\n')

par4 = doc.add_paragraph(f'2. Итого по Договору в Отчетном периоде оказано Услуг на сумму 570 000,00 руб. '
                     '(Пятьсот семьдесят тысяч рублей 00 копеек) кроме того НДС 114 000,оо руб, исчисляемый в '
                     'соответствии со статьей 164 НК РФ')
par4.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
fmt = par4.paragraph_format
fmt.first_line_indent = Mm(12.7)

par5 = doc.add_paragraph('3. Перечисленные Услуги оказаны своевременно в необходимом объеме и в '
                         'соответствии с требованиями, установленными Договором к качеству услуг. ')
par5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
fmt = par5.paragraph_format
fmt.first_line_indent = Mm(12.7)

par6 = doc.add_paragraph('4. Акт составлен и подписан в двух экземплярах, имеющих равную силу, по одному '
                         'экземпляру для каждой Стороны. ')
par6.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
fmt = par6.paragraph_format
fmt.first_line_indent = Mm(12.7)

tail = doc.add_paragraph('')
tail.add_run('Подписи сторон').bold = True
tail.alignment = WD_ALIGN_PARAGRAPH.CENTER

tail_table = doc.add_table(2, 2)
tail_table.style = 'Table Grid'
cell = tail_table.cell(0, 0).add_paragraph('')
cell.add_run(f"ЗАКАЗЧИК\nИсполнительный директор\nАО «МР Групп»").bold = True
cell = tail_table.cell(0, 1).add_paragraph('')
cell.add_run(f"ИСПОЛНИТЕЛЬ\nООО «Мультиспейс Павелецкая»\nГенеральный директор\nуправляющей организации \nООО «Прайдекс Констракшн»").bold = True
cell = tail_table.cell(1, 0).add_paragraph('')
cell.add_run(f"_____________/Обухов А.А./\n").bold = True
cell = tail_table.cell(1, 1).add_paragraph('')
cell.add_run(f"____________/ Бренев К.С./\n").bold = True




# Сохранение документа
doc.save(f'C:\\Users\\anton.kudryashov\\Documents\\Акт сдачи-приемки оказанных услуг (Тех поддержка) {report_month:02}.22.docx')
